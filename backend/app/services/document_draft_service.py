from __future__ import annotations

import base64
import io
import json
import re
from typing import Any

from docx import Document
from docx.shared import Inches


def _safe_str(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def dumps_content(content: dict[str, Any] | None) -> str:
    return json.dumps(content or {}, ensure_ascii=False)


def loads_content(raw: str | None) -> dict[str, Any]:
    if not raw:
        return {}
    try:
        value = json.loads(raw)
        if isinstance(value, dict):
            return value
        return {}
    except Exception:
        return {}


def render_application_markdown(content: dict[str, Any]) -> str:
    lines: list[str] = []
    lines.append("# 申请书（Markdown 预览）")
    lines.append("")
    lines.append("## 基本信息")
    lines.append("")
    lines.append(f"- 项目编号：{_safe_str(content.get('project_number')) or '-'}")
    lines.append(f"- 项目名称：{_safe_str(content.get('project_name')) or '-'}")
    lines.append(f"- 团队负责人：{_safe_str(content.get('leader_name')) or '-'}")
    lines.append(f"- 手机号码：{_safe_str(content.get('leader_phone')) or '-'}")
    lines.append(
        f"- 项目时间：{(_safe_str(content.get('start_date')) + ' 至 ' + _safe_str(content.get('end_date'))).strip(' 至') or '-'}"
    )
    lines.append("")

    tech_points = content.get("tech_points") or []
    if isinstance(tech_points, list):
        tech_list = [x for x in (_safe_str(t) for t in tech_points) if x]
    else:
        tech_list = [x for x in [_safe_str(tech_points)] if x]

    lines.append("## 技术点")
    lines.append("")
    if tech_list:
        for t in tech_list:
            lines.append(f"- {t}")
    else:
        lines.append("- -")
    lines.append("")

    participants = content.get("participants") or []
    lines.append("## 参加人员")
    lines.append("")
    if isinstance(participants, list) and participants:
        for idx, p in enumerate(participants, start=1):
            if not isinstance(p, dict):
                continue
            name = _safe_str(p.get("name")) or "-"
            sid = _safe_str(p.get("student_id")) or "-"
            major = _safe_str(p.get("major")) or "-"
            phone = _safe_str(p.get("phone")) or "-"
            email = _safe_str(p.get("email")) or "-"
            role = _safe_str(p.get("role")) or "-"
            lines.append(f"{idx}. {name}（{role}）")
            lines.append(f"   - 学号：{sid}")
            lines.append(f"   - 专业：{major}")
            lines.append(f"   - 电话：{phone}")
            lines.append(f"   - 邮箱：{email}")
    else:
        lines.append("- -")
    lines.append("")

    other_people = _safe_str(content.get("other_people"))
    lines.append("## 其他人员")
    lines.append("")
    lines.append(other_people or "-")
    lines.append("")

    project_blog_url = _safe_str(content.get("project_blog_url"))
    member_blog_urls = content.get("member_blog_urls") or []
    lines.append("## 博客与资料")
    lines.append("")
    lines.append(f"- 项目资料地址：{project_blog_url or '-'}")
    if isinstance(member_blog_urls, list):
        urls = [u for u in (_safe_str(x) for x in member_blog_urls) if u]
    else:
        urls = [u for u in [_safe_str(member_blog_urls)] if u]
    if urls:
        lines.append("- 成员博客：")
        for u in urls:
            lines.append(f"  - {u}")
    else:
        lines.append("- 成员博客：-")
    lines.append("")

    lines.append("## 项目介绍")
    lines.append("")
    lines.append(_safe_str(content.get("project_intro")) or "-")
    lines.append("")

    lines.append("## 项目目标")
    lines.append("")
    lines.append(_safe_str(content.get("project_goal")) or "-")
    lines.append("")

    lines.append("## 实施计划")
    lines.append("")
    lines.append(_safe_str(content.get("plan")) or "-")
    lines.append("")

    lines.append("## 预期成果")
    lines.append("")
    lines.append(_safe_str(content.get("expected_results")) or "-")
    lines.append("")

    lines.append("## 指导教师意见")
    lines.append("")
    lines.append(_safe_str(content.get("teacher_comment")) or "-")
    lines.append("")

    lines.append("## 承诺日期")
    lines.append("")
    lines.append(_safe_str(content.get("commitment_date")) or "-")
    lines.append("")

    return "\n".join(lines).strip() + "\n"


def render_task_markdown(content: dict[str, Any]) -> str:
    lines: list[str] = []
    lines.append("# 任务书（Markdown 预览）")
    lines.append("")
    lines.append("## 基本信息")
    lines.append("")
    lines.append(f"- 项目名称：{_safe_str(content.get('project_name')) or '-'}")
    lines.append(f"- 团队名称：{_safe_str(content.get('team_name')) or '-'}")
    lines.append(f"- 负责人：{_safe_str(content.get('leader_name')) or '-'}")
    lines.append(f"- 联系方式：{_safe_str(content.get('leader_phone')) or '-'}")
    lines.append(
        f"- 实施时间：{(_safe_str(content.get('start_date')) + ' 至 ' + _safe_str(content.get('end_date'))).strip(' 至') or '-'}"
    )
    lines.append("")

    lines.append("## 任务目标")
    lines.append("")
    lines.append(_safe_str(content.get("objectives")) or "-")
    lines.append("")

    roles = content.get("roles") or []
    lines.append("## 任务分解与分工")
    lines.append("")
    if isinstance(roles, list) and roles:
        for idx, r in enumerate(roles, start=1):
            if not isinstance(r, dict):
                continue
            name = _safe_str(r.get("name")) or "-"
            responsibility = _safe_str(r.get("responsibility")) or "-"
            deliverable = _safe_str(r.get("deliverable")) or "-"
            lines.append(f"{idx}. {name}")
            lines.append(f"   - 职责：{responsibility}")
            lines.append(f"   - 交付物：{deliverable}")
    else:
        lines.append(_safe_str(content.get("role_text")) or "-")
    lines.append("")

    milestones = content.get("milestones") or []
    lines.append("## 阶段计划")
    lines.append("")
    if isinstance(milestones, list) and milestones:
        for idx, m in enumerate(milestones, start=1):
            if not isinstance(m, dict):
                continue
            phase = _safe_str(m.get("phase")) or "-"
            start = _safe_str(m.get("start")) or "-"
            end = _safe_str(m.get("end")) or "-"
            work = _safe_str(m.get("work")) or "-"
            deliverable = _safe_str(m.get("deliverable")) or "-"
            lines.append(f"{idx}. {phase}（{start} - {end}）")
            lines.append(f"   - 工作内容：{work}")
            lines.append(f"   - 阶段产出：{deliverable}")
    else:
        lines.append(_safe_str(content.get("plan")) or "-")
    lines.append("")

    lines.append("## 风险与应对")
    lines.append("")
    lines.append(_safe_str(content.get("risks")) or "-")
    lines.append("")

    lines.append("## 验收标准")
    lines.append("")
    lines.append(_safe_str(content.get("acceptance")) or "-")
    lines.append("")

    return "\n".join(lines).strip() + "\n"


def _add_kv_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for idx, (k, v) in enumerate(rows):
        table.cell(idx, 0).text = k
        table.cell(idx, 1).text = v


def _data_url_to_image_bytes(data_url: Any) -> tuple[bytes, str] | None:
    if not isinstance(data_url, str):
        return None
    raw = data_url.strip()
    if not raw.startswith("data:"):
        return None

    m = re.match(r"^data:(?P<mime>[^;]+);base64,(?P<b64>.+)$", raw, flags=re.IGNORECASE | re.DOTALL)
    if not m:
        return None
    mime = (m.group("mime") or "").strip().lower()
    b64 = (m.group("b64") or "").strip()
    if not b64:
        return None
    try:
        data = base64.b64decode(b64)
    except Exception:
        return None
    if not data:
        return None
    return data, mime


def render_application_docx(content: dict[str, Any]) -> bytes:
    doc = Document()

    doc.add_paragraph("软件学院创新项目实训")
    doc.add_paragraph("申请表")
    doc.add_paragraph("（2026版）")
    doc.add_paragraph("")

    _add_kv_table(
        doc,
        [
            ("项目编号（申请人不必填写）", _safe_str(content.get("project_number"))),
            ("项目名称", _safe_str(content.get("project_name"))),
            ("团队负责人", _safe_str(content.get("leader_name"))),
            ("手机号码", _safe_str(content.get("leader_phone"))),
            (
                "项目时间",
                f"{_safe_str(content.get('start_date'))} 至 {_safe_str(content.get('end_date'))}".strip(),
            ),
        ],
    )

    doc.add_paragraph("")
    doc.add_paragraph("山东大学软件学院")
    doc.add_page_break()

    doc.add_paragraph("填报说明")
    instructions = [
        "1.项目目标请简短说明项目的研发定位。",
        "2.指导教师不需要填写，如果有导师指导，立项通过后说明。",
        "3.技术点请列出项目中使用的关键技术点，最多不超过5个名词。",
        "4.参加人员中的第一行必须为项目负责人。团队总人数3-4人，少于3人、超过4人（小于等于2人，或者大于等于5人）需要提前发邮件申请。所有成员必须为2023级软件学院的大三在校本科生，其他参与人员请在其他人员中列出。",
        "5.项目资料地址必须可公开访问，项目组每个成员都需要维护各自的博客，以便督导老师随时可以抽查项目实施过程。必须选用新浪、CSDN等知名服务提供商，不支持自己搭建的服务器。",
        "6.项目介绍从项目背景、技术创新、工作内容、技术路线、实施方案等方面，详细说明项目的具体工作。请详细说明，可以添加附件，或者音视频等多媒体材料的网址。",
        "7.实施计划按照时间节点，分阶段完成预定目标。",
        "8.申请承诺需要团队所有成员签字。",
    ]
    for line in instructions:
        doc.add_paragraph(line)

    doc.add_paragraph("")

    tech_points = content.get("tech_points") or []
    if isinstance(tech_points, list):
        tech = "、".join([_safe_str(x) for x in tech_points if _safe_str(x)])[:500]
    else:
        tech = _safe_str(tech_points)

    _add_kv_table(
        doc,
        [
            ("项目名称", _safe_str(content.get("project_name"))),
            (
                "实施时间",
                f"{_safe_str(content.get('start_date'))} 至 {_safe_str(content.get('end_date'))}".strip(),
            ),
            ("项目目标", _safe_str(content.get("project_goal"))),
            ("技术要点", tech),
            ("指导教师", _safe_str(content.get("teacher"))),
            ("团队名称", _safe_str(content.get("team_name"))),
        ],
    )

    doc.add_paragraph("")
    doc.add_paragraph("参加人员")
    participants = content.get("participants") or []
    if not isinstance(participants, list):
        participants = []

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["姓名", "学号", "专业", "手机号码", "电子邮箱", "项目分工", "电子签名"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for p in participants:
        if not isinstance(p, dict):
            continue
        row = table.add_row().cells
        row[0].text = _safe_str(p.get("name"))
        row[1].text = _safe_str(p.get("student_id"))
        row[2].text = _safe_str(p.get("major"))
        row[3].text = _safe_str(p.get("phone"))
        row[4].text = _safe_str(p.get("email"))
        row[5].text = _safe_str(p.get("role"))
        image = _data_url_to_image_bytes(p.get("signature_data_url"))
        if image:
            data, _mime = image
            try:
                run = row[6].paragraphs[0].add_run()
                run.add_picture(io.BytesIO(data), width=Inches(1.6))
            except Exception:
                row[6].text = ""
        else:
            row[6].text = ""

    doc.add_paragraph("")
    doc.add_paragraph("其他人员")
    doc.add_paragraph(_safe_str(content.get("other_people")))

    doc.add_paragraph("")
    doc.add_paragraph("项目资料地址")
    doc.add_paragraph(f"项目博客地址：{_safe_str(content.get('project_blog_url'))}")
    member_blogs = content.get("member_blog_urls") or []
    if not isinstance(member_blogs, list):
        member_blogs = []
    doc.add_paragraph("成员个人博客地址：")
    for url in [u for u in member_blogs if _safe_str(u)]:
        doc.add_paragraph(_safe_str(url))

    doc.add_paragraph("")
    doc.add_paragraph("项目介绍（可跨页，可附加多媒体网址）")
    doc.add_paragraph(_safe_str(content.get("project_intro")))

    doc.add_paragraph("")
    doc.add_paragraph("实施计划")
    doc.add_paragraph(_safe_str(content.get("plan")))

    doc.add_paragraph("")
    doc.add_paragraph("预期成果")
    doc.add_paragraph(_safe_str(content.get("expected_results")))

    doc.add_paragraph("")
    doc.add_paragraph("指导教师评语")
    doc.add_paragraph(_safe_str(content.get("teacher_comment")))

    doc.add_paragraph("")
    doc.add_paragraph("本人郑重承诺，此申请书内容真实有效。")
    doc.add_paragraph("（所有团队成员签字）")
    doc.add_paragraph(_safe_str(content.get("commitment_date")))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_task_docx(content: dict[str, Any]) -> bytes:
    doc = Document()

    doc.add_paragraph("软件学院创新项目实训")
    doc.add_paragraph("任务书")
    doc.add_paragraph("（2026版）")
    doc.add_paragraph("")

    _add_kv_table(
        doc,
        [
            ("项目名称", _safe_str(content.get("project_name"))),
            ("实施时间", f"{_safe_str(content.get('start_date'))} 至 {_safe_str(content.get('end_date'))}".strip()),
            ("团队名称", _safe_str(content.get("team_name"))),
            ("负责人", _safe_str(content.get("leader_name"))),
            ("联系方式", _safe_str(content.get("leader_phone"))),
        ],
    )

    doc.add_paragraph("")
    doc.add_paragraph("任务目标")
    doc.add_paragraph(_safe_str(content.get("objectives")))

    doc.add_paragraph("")
    doc.add_paragraph("任务分解与分工")
    roles = content.get("roles") or []
    if isinstance(roles, list) and roles:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.cell(0, 0).text = "成员"
        table.cell(0, 1).text = "职责"
        table.cell(0, 2).text = "交付物"
        for r in roles:
            if not isinstance(r, dict):
                continue
            row = table.add_row().cells
            row[0].text = _safe_str(r.get("name"))
            row[1].text = _safe_str(r.get("responsibility"))
            row[2].text = _safe_str(r.get("deliverable"))
    else:
        doc.add_paragraph(_safe_str(content.get("role_text")))

    doc.add_paragraph("")
    doc.add_paragraph("阶段计划")
    milestones = content.get("milestones") or []
    if isinstance(milestones, list) and milestones:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.cell(0, 0).text = "阶段"
        table.cell(0, 1).text = "起止时间"
        table.cell(0, 2).text = "工作内容"
        table.cell(0, 3).text = "阶段产出"
        for m in milestones:
            if not isinstance(m, dict):
                continue
            row = table.add_row().cells
            row[0].text = _safe_str(m.get("phase"))
            row[1].text = f"{_safe_str(m.get('start'))} - {_safe_str(m.get('end'))}".strip()
            row[2].text = _safe_str(m.get("work"))
            row[3].text = _safe_str(m.get("deliverable"))
    else:
        doc.add_paragraph(_safe_str(content.get("plan")))

    doc.add_paragraph("")
    doc.add_paragraph("风险与应对")
    doc.add_paragraph(_safe_str(content.get("risks")))

    doc.add_paragraph("")
    doc.add_paragraph("验收标准")
    doc.add_paragraph(_safe_str(content.get("acceptance")))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
